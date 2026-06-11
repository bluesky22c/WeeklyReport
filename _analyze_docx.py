import os

from docx import Document

path = r"D:\Temp\주간업무정리\[보고_제어1팀]주간업무보고(260529).docx"
out = []
doc = Document(path)

out.append("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        out.append(f"[P{i}] style={p.style.name} | {p.text}")

out.append("")
out.append("=== TABLES ===")
for ti, t in enumerate(doc.tables):
    out.append(f"--- TABLE {ti}: {len(t.rows)} rows x {len(t.columns)} cols ---")
    for ri, row in enumerate(t.rows):
        cells = [c.text.replace("\n", " / ").strip() for c in row.cells]
        out.append(f"  R{ri}: {cells}")

out.append("")
out.append("=== HEADERS / FOOTERS ===")
for si, sec in enumerate(doc.sections):
    for p in sec.header.paragraphs:
        if p.text.strip():
            out.append(f"[HDR s{si}] {p.text}")
    for p in sec.footer.paragraphs:
        if p.text.strip():
            out.append(f"[FTR s{si}] {p.text}")

with open(r"D:\Temp\주간업무정리\_docx_struct.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(out))
print(f"done; paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}")
