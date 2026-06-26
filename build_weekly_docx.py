from __future__ import annotations

import os
import re
import zipfile
from collections import OrderedDict
from copy import deepcopy
from datetime import datetime

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BASE = os.path.dirname(os.path.abspath(__file__))
PEOPLE = ["박주상", "주영모", "김병걸", "김현석", "송유림", "김성현", "배유민", "박장미", "김도현", "유찬", "황지환"]
THIS_WEEK_LABEL = "2026-06-22 ~ 2026-06-26"
NEXT_WEEK_LABEL = "2026-06-29 ~ 2026-07-03"


def clean_text(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n\s*\n+", "\n", value)
    return value.strip()


def cell_text(cell) -> str:
    parts = []
    for p in cell.find_all("p"):
        text = clean_text(p.get_text("\n", strip=True))
        if text:
            parts.append(text)
    if parts:
        return "\n".join(parts)
    return clean_text(cell.get_text("\n", strip=True))


def latest_report_path(name: str) -> str | None:
    folder = os.path.join(BASE, name)
    if not os.path.isdir(folder):
        return None
    files = sorted(f for f in os.listdir(folder) if re.fullmatch(r"2026-06-2[56]\.html", f))
    if not files:
        return None
    return os.path.join(folder, files[-1])


def previous_template_path(current_yymmdd: str) -> str:
    candidates: list[tuple[str, str]] = []
    for fname in os.listdir(BASE):
        m = re.fullmatch(r"\[보고_제어1팀\]주간업무보고\((\d{6})\)\.docx", fname)
        if m and m.group(1) < current_yymmdd:
            candidates.append((m.group(1), os.path.join(BASE, fname)))
    if not candidates:
        raise RuntimeError("전주 통합 DOCX 템플릿을 찾지 못했습니다.")
    return sorted(candidates)[-1][1]


def read_rows(path: str, fallback_owner: str) -> tuple[str, list[list[str]]]:
    with open(path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    source_el = soup.select_one("p.source")
    source = clean_text(source_el.get_text(" ", strip=True)) if source_el else ""
    table = None
    for candidate in soup.find_all("table"):
        text = clean_text(candidate.get_text(" ", strip=True))
        if "프로젝트" in text and ("금주" in text or "계획" in text) and "내주" in text:
            table = candidate
    if table is None:
        return source, []

    rows: list[list[str]] = []
    pending_project_prefix = ""
    for tr in table.find_all("tr"):
        cells = [cell_text(c) for c in tr.find_all(["th", "td"], recursive=False)]
        if not cells or not any(cells):
            continue
        if cells[0].replace(" ", "") == "프로젝트":
            continue
        if len(cells) == 1:
            text = cells[0].strip()
            if text and text not in {"기타"}:
                pending_project_prefix = text
            continue
        if len(cells) == 2 and rows:
            if cells[0].strip():
                rows[-1][1] = clean_text(f"{rows[-1][1]}\n{cells[0]}")
            if cells[1].strip():
                rows[-1][3] = clean_text(f"{rows[-1][3]}\n{cells[1]}")
            continue
        while len(cells) < 6:
            cells.append("")
        project, this_week, owner1, next_week, owner2, target = cells[:6]
        if pending_project_prefix:
            project = clean_text(f"{pending_project_prefix}\n{project}")
            pending_project_prefix = ""
        owner = owner1 or owner2 or fallback_owner
        if not any([project, this_week, owner, next_week, target]):
            continue
        rows.append([project, this_week, owner, next_week, target])
    return source, rows


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size_pt: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def write_cell(cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    lines = text.splitlines() or [""]
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        set_run_font(run, 8, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def compact(value: str) -> str:
    return re.sub(r"\s+", "", value).lower()


def classify_docx_row(member: str, row: list[str]) -> tuple[str, str]:
    project, this_week = row[0], row[1]
    raw = f"{project}\n{this_week}"
    raw_c = compact(raw)

    if any(k in raw for k in ["Multi-Probing", "MPT", "EPI-SE1000"]):
        return "SEC", "L-TF Multi-Probing Tester (MPT) 고객대응"
    if any(k in raw for k in ["신뢰성", "LEDoS", "CSS", "Vision Part"]):
        if member in {"김병걸", "김도현"} or "신뢰성" in raw or "LEDoS" in raw:
            return "SEC", "CSS LEDoS 신뢰성 평가 설비 Setup"
    if any(k in raw for k in ["AR System", "Sparse", "AF", "AM35", "AR "]):
        return "SEC", "AR System"

    if any(k in raw for k in ["SFA Micro LED-N", "전계인가", "LCR Meter", "IM3533"]):
        return "SDC", "SFA Micro LED-N SWP 설비용 전계인가시스템"
    if any(k in raw for k in ["SDC A3", "LifeTimeSystem", "Pallet"]):
        return "SDC", "SDC A3 LifeTimeSystem"

    if any(k in raw for k in ["ICBSystem", "LGD"]):
        if "ToE" not in raw and "TOE" not in raw and "CIM" not in raw:
            return "LGD", "ICB System"
    if "toe" in raw_c or "cim" in raw_c:
        if "CIM" in raw or "cim" in raw_c or "2호기" in raw:
            return "LGD", "ToE Full Contact 터치 검사 설비 Setup 2호기"
    if ("toe" in raw_c or "toe" in raw.lower()) and ("1호기" in raw or "1 호기" in raw or "Himax" in raw):
        return "LGD", "ToE Full Contact 터치 검사 설비 Setup 1호기"
    if "toe" in raw_c:
        return "LGD", "ToE Full Contact 터치 검사 설비 Setup 2호기"

    if any(k in raw for k in ["WHTM", "ACAControlSystem"]):
        return "WHTM", "G6 ACAControlSystem"

    if any(k in raw for k in ["Micro LED", "MicroLED", "WPF Framework", "KOPTI (광기술원)"]):
        if member in {"박주상", "유찬"}:
            return "KOPTI / CTP", "Micro LED Inspector (EMI-K026)"
    if any(k in raw for k in ["광명 광기술원", "Sorter", "Prober"]):
        return "KOPTI / CTP", "KOPTI 설비 이전 대응"
    if "CTP" in raw:
        return "KOPTI / CTP", "CTP"

    return "기타", project or "기타"


def split_mixed_toe_row(row: list[str]) -> list[list[str]]:
    project, this_week, owner, next_week, target = row
    raw = f"{project}\n{this_week}\n{next_week}"
    if not ("TOE" in raw or "ToE" in raw) or "1호기" not in raw or "CIM" not in raw:
        return [row]

    def split_text(value: str) -> tuple[str, str]:
        markers = ["1.TOE 1호기", "1. TOE 1호기", "1.TOE  1호기", "1.ToE 1호기", "1. ToE 1호기"]
        for marker in markers:
            idx = value.find(marker)
            if idx >= 0:
                return clean_text(value[:idx]), clean_text(value[idx:])
        return value, ""

    this_week_2, this_week_1 = split_text(this_week)
    next_week_2, next_week_1 = split_text(next_week)
    rows = []
    if this_week_1 or next_week_1:
        rows.append(["ToE Full Contact 터치 검사 설비 Setup 1호기", this_week_1, owner, next_week_1, target])
    if this_week_2 or next_week_2:
        rows.append(["ToE Full Contact 터치 검사 설비 Setup 2호기", this_week_2, owner, next_week_2, target])
    return rows or [row]


def append_unique_lines(existing: str, value: str) -> str:
    parts = [p.strip() for p in existing.split("\n\n") if p.strip()]
    value = value.strip()
    if not value:
        return existing.strip()
    if value not in parts:
        parts.append(value)
    return "\n\n".join(parts)


def merge_owner(existing: str, owner: str) -> str:
    names = []
    for chunk in re.split(r"[\n/,]+", existing):
        chunk = chunk.strip()
        if chunk and chunk not in names:
            names.append(chunk)
    for chunk in re.split(r"[\n/,]+", owner):
        chunk = chunk.strip()
        if chunk and chunk not in names:
            names.append(chunk)
    return "\n".join(names)


def build_grouped_rows(reports: list[dict]) -> list[tuple[str, list[str]]]:
    grouped: OrderedDict[tuple[str, str], list[str]] = OrderedDict()
    group_order = ["SEC", "SDC", "LGD", "WHTM", "KOPTI / CTP", "기타"]
    title_order = {
        "AR System": 10,
        "L-TF Multi-Probing Tester (MPT) 고객대응": 20,
        "CSS LEDoS 신뢰성 평가 설비 Setup": 30,
        "SDC A3 LifeTimeSystem": 10,
        "SFA Micro LED-N SWP 설비용 전계인가시스템": 20,
        "ICB System": 10,
        "ToE Full Contact 터치 검사 설비 Setup 1호기": 20,
        "ToE Full Contact 터치 검사 설비 Setup 2호기": 30,
        "G6 ACAControlSystem": 10,
        "Micro LED Inspector (EMI-K026)": 10,
        "KOPTI 설비 이전 대응": 20,
        "CTP": 30,
    }

    for report in reports:
        member = report["name"]
        for source_row in report["rows"]:
            for normalized_row in split_mixed_toe_row(source_row):
                group, title = classify_docx_row(member, normalized_row)
                key = (group, title)
                if key not in grouped:
                    grouped[key] = [title, "", "", "", ""]
                row = grouped[key]
                row[1] = append_unique_lines(row[1], normalized_row[1])
                row[2] = merge_owner(row[2], normalized_row[2] or member)
                row[3] = append_unique_lines(row[3], normalized_row[3])
                row[4] = append_unique_lines(row[4], normalized_row[4])

    output: list[tuple[str, list[str]]] = []
    for group in group_order:
        rows = [(g, values) for (g, _title), values in grouped.items() if g == group]
        rows.sort(key=lambda item: (title_order.get(item[1][0], 999), item[1][0]))
        if not rows:
            continue
        output.append(("group", [group] * 5))
        for _g, values in rows:
            output.append(("data", values))
    return output


def set_table_header_text(table, latest_date: str) -> None:
    headers = [
        "프로젝트",
        f"금주 핵심 실적 ({THIS_WEEK_LABEL})",
        "담당자",
        f"내주 계획 ({NEXT_WEEK_LABEL})",
        "목표일",
    ]
    for idx, text in enumerate(headers):
        write_cell(table.rows[0].cells[idx], text, bold=True, center=True)


def clear_table_body(table) -> None:
    tbl = table._tbl
    for tr in list(tbl.tr_lst)[1:]:
        tbl.remove(tr)


def append_template_row(table, template_tr, values: list[str], *, row_type: str) -> None:
    table._tbl.append(deepcopy(template_tr))
    row = table.rows[-1]
    for idx, value in enumerate(values):
        write_cell(row.cells[idx], value, bold=(row_type == "group"), center=(row_type == "group" or idx in (2, 4)))


def build_docx() -> str:
    reports: list[dict] = []
    missing: list[str] = []
    for name in PEOPLE:
        path = latest_report_path(name)
        if not path:
            missing.append(name)
            continue
        source, rows = read_rows(path, name)
        reports.append({
            "name": name,
            "date": os.path.splitext(os.path.basename(path))[0],
            "path": path,
            "source": source,
            "rows": rows,
        })

    if not reports:
        raise RuntimeError("이번 주 HTML 보고서를 찾지 못했습니다.")

    latest_date = max(r["date"] for r in reports)
    yymmdd = datetime.strptime(latest_date, "%Y-%m-%d").strftime("%y%m%d")
    out_path = os.path.join(BASE, f"[보고_제어1팀]주간업무보고({yymmdd}).docx")
    template_path = previous_template_path(yymmdd)

    doc = Document(template_path)

    styles = doc.styles
    styles["Normal"].font.name = "맑은 고딕"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    styles["Normal"].font.size = Pt(8)

    title = doc.paragraphs[0]
    title.text = ""
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("주간업무 실적 및 계획 (제어1팀)")
    set_run_font(title_run, 12, bold=True)

    meta = doc.paragraphs[1]
    meta.text = ""
    meta_run = meta.add_run(f"Date : {latest_date}     작성자 : 이종규")
    set_run_font(meta_run, 8)

    table = doc.tables[0]
    group_template = deepcopy(table.rows[1]._tr)
    data_template = deepcopy(table.rows[2]._tr)
    set_table_header_text(table, latest_date)
    clear_table_body(table)

    for row_type, values in build_grouped_rows(reports):
        append_template_row(table, group_template if row_type == "group" else data_template, values, row_type=row_type)

    doc.save(out_path)
    with zipfile.ZipFile(out_path) as zf:
        bad = zf.testzip()
        if bad:
            raise RuntimeError(f"DOCX ZIP 검사 실패: {bad}")
    print(f"created={out_path}")
    print(f"latest_date={latest_date}")
    print(f"reports={len(reports)}")
    print(f"source_rows={sum(len(r['rows']) for r in reports)}")
    print(f"docx_rows={len(build_grouped_rows(reports)) + 1}")
    print(f"missing={','.join(missing)}")
    print(f"template={template_path}")
    return out_path


if __name__ == "__main__":
    build_docx()
